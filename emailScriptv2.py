import xlrd
import docx
import os
from num2words import num2words

path = "F:\\SOSLAVINO\\sosl.xls"

data = xlrd.open_workbook(path)
data_sheet = data.sheet_by_index(0)

housenumber = []
print(data_sheet.nrows)
rowmax = data_sheet.nrows

for row in range(1, rowmax):
    housenumber.append(int(data_sheet.cell_value(row, 0)))

name = []
for row in range(1, rowmax):
    name.append(data_sheet.cell_value(row, 1))

adress = []
for row in range(1, rowmax):
    adress.append(data_sheet.cell_value(row, 2))

cadastral = []
for row in range(1, rowmax):
    cadastral.append(data_sheet.cell_value(row, 3))

dept = []
for row in range(1, rowmax):
    if data_sheet.cell_value(row, 5) != '':
        dept.append(int(data_sheet.cell_value(row, 5)))
    else:
        dept.append(0)

dept2 = []
for row in range(1, rowmax):
    if data_sheet.cell_value(row, 4) != '':
        dept2.append(int(data_sheet.cell_value(row, 4)))
    else:
        dept2.append(0)

target = []
for row in range(1, rowmax):
    if data_sheet.cell_value(row, 6) != '':
        target.append(0)
    else:
        target.append(1000)

doc = docx.Document("F:\\SOSLAVINO\\notification.docx")
os.chdir("F:\\SOSLAVINO\\notif0912")

for i in range(len(housenumber)):
    doc.paragraphs[7].clear()
    doc.paragraphs[7].add_run(name[i])
    doc.paragraphs[9].clear()
    doc.paragraphs[9].add_run(adress[i])
    doc.paragraphs[19].clear()
    new_name = name[i].split()
    doc.paragraphs[19].add_run("Уважаемый(ая) " + new_name[1] + ' ' + new_name[2] + '!')
    doc.paragraphs[22].clear()
    doc.paragraphs[22].add_run("Вы являетесь собственником земельного участка № " + str(housenumber[i]) + ' ' + "с кадастровым номером: " + str(cadastral[i]) + " (1000 кв. м.) расположенного в границах Садоводческого некоммерческого товарищества   «Сославино» по адресу: Московская область, Волоколамский район, Болычевский с.о., район дер. Дьяково и членом Садоводческого некоммерческого товарищества   «Сославино». ")
    doc.paragraphs[28].clear()
    doc.paragraphs[28].add_run("Общая сумма Вашей задолженности по членским и целевым взносам, за период с 2018 года по 28 ноября 2021 года, без учета пени, составляет " + str(dept[i]) + " (" + num2words(dept[i], lang = 'rus') + ") рублей.")
    if target[i] != 0:
        doc.paragraphs[28].add_run(" Целевой взнос на вырубку под ЛЭП - 1000 (одна тысяча) рублей.")
        #doc.paragraphs[40].clear()
        #doc.paragraphs[40].add_run("На основании вышеизложенного и руководствуясь действующим законодательством Российской Федерации,  просим Вас погасить имеющуюся задолженность в размере " + str(dept[i] + dept2[i] + target[i]) + " (" + num2words(dept[i] + dept2[i] + target[i], lang = 'rus') + ") рублей в срок до: «15» нояября 2021 года.")
    else:
        pass
        #doc.paragraphs[40].clear()
        #doc.paragraphs[40].add_run("На основании вышеизложенного и руководствуясь действующим законодательством Российской Федерации,  просим Вас погасить имеющуюся задолженность в размере " + str(dept[i]) + " (" + num2words(dept[i] + dept2[i], lang = 'rus') + ") рублей в срок до: «15» ноября 2021 года.")       
    if dept2[i] != 0:
        doc.paragraphs[28].add_run(" Паевой взнос на строительство ЛЭП " + str(dept2[i]) + " (" + num2words(dept2[i], lang = 'rus') + ") рублей.")
        #doc.paragraphs[40].clear()
        #doc.paragraphs[40].add_run("На основании вышеизложенного и руководствуясь действующим законодательством Российской Федерации,  просим Вас погасить имеющуюся задолженность в размере " + str(dept[i] + dept2[i] + target[i]) + " (" + num2words(dept[i] + dept2[i] + target[i], lang = 'rus') + ") рублей в срок до: «15» ноября 2021 года.") 
        #doc.paragraphs[40].clear()
        #doc.paragraphs[40].add_run("На основании вышеизложенного и руководствуясь действующим законодательством Российской Федерации,  просим Вас погасить имеющуюся задолженность в размере " + str(dept[i] + dept2[i] + 1000) + " (" + num2words(dept[i] + dept2[i] + 1000, lang = 'rus') + ") рублей в срок до: «15» ноября 2021 года.")
    doc.paragraphs[39].clear()
    doc.paragraphs[39].add_run("На основании вышеизложенного и руководствуясь действующим законодательством Российской Федерации,  просим Вас погасить имеющуюся задолженность в размере " + str(dept[i] + dept2[i] + target[i]) + " (" + num2words(dept[i] + dept2[i] + target[i], lang = 'rus') + ") рублей в срок до: «15» декабря 2021 года.")
    doc.save(name[i] + " участок " + str(housenumber[i]) + '.docx')

#print(type(dept[9]))